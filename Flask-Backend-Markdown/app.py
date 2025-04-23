from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import tempfile
import uuid
from docx import Document
import html2text
import io
import subprocess
import platform

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

TEMP_FOLDER = tempfile.gettempdir()
if not os.path.exists(TEMP_FOLDER):
    os.makedirs(TEMP_FOLDER)

@app.route('/merge-and-convert', methods=['POST'])
def merge_and_convert():
    try:
        if 'files[]' not in request.files:
            return jsonify({"error": "No files provided"}), 400
        
        files = request.files.getlist('files[]')
        output_format = request.form.get('format', 'markdown')
        
        if not files or files[0].filename == '':
            return jsonify({"error": "No files selected"}), 400
        
        processed_files = []
        for file in files:
            if file.filename.endswith('.docx'):
                temp_filename = os.path.join(TEMP_FOLDER, f"{uuid.uuid4()}.docx")
                file.save(temp_filename)
                processed_files.append(temp_filename)
            elif file.filename.endswith('.doc'):
                temp_doc = os.path.join(TEMP_FOLDER, f"{uuid.uuid4()}.doc")
                file.save(temp_doc)
                
                temp_docx = os.path.join(TEMP_FOLDER, f"{uuid.uuid4()}.docx")
                doc_to_docx(temp_doc, temp_docx)
                
                os.remove(temp_doc)
                processed_files.append(temp_docx)
            else:
                return jsonify({"error": f"File {file.filename} is not a DOC or DOCX file"}), 400
        
        merged_content = merge_docx_files(processed_files)
        
        if output_format == 'markdown':
            output_content, output_filename, output_mimetype = convert_to_markdown(merged_content)
        else:
            output_content, output_filename, output_mimetype = convert_to_text(merged_content)
        
        return send_file(
            io.BytesIO(output_content),
            as_attachment=True,
            download_name=output_filename,
            mimetype=output_mimetype
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def doc_to_docx(doc_path, docx_path):
    system = platform.system()
    
    try:
        if system == 'Windows':
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.visible = False
            doc = word.Documents.Open(os.path.abspath(doc_path))
            doc.SaveAs(os.path.abspath(docx_path), 16)
            doc.Close()
            word.Quit()
        else:
            # Try common paths for LibreOffice executable
            libreoffice_paths = [
                'libreoffice',
                '/usr/bin/libreoffice',
                '/usr/local/bin/libreoffice',
                '/opt/libreoffice/program/soffice'
            ]
            
            cmd = None
            for path in libreoffice_paths:
                try:
                    # Test if the command exists
                    subprocess.run([path, '--version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
                    cmd = [
                        path, '--headless', '--convert-to', 'docx',
                        '--outdir', os.path.dirname(docx_path), doc_path
                    ]
                    break
                except (subprocess.SubprocessError, FileNotFoundError):
                    continue
            
            if cmd is None:
                raise Exception("LibreOffice dont install. Please install by command: sudo apt-get install libreoffice")
                
            subprocess.run(cmd, check=True)
            base_name = os.path.basename(doc_path)
            original_output = os.path.join(os.path.dirname(docx_path), os.path.splitext(base_name)[0] + '.docx')
            os.rename(original_output, docx_path)
    except Exception as e:
        raise Exception(f"Error converting DOC to DOCX: {str(e)}")

def merge_docx_files(files):
    merged_doc = Document()
    
    for i, file_path in enumerate(files):
        doc = Document(file_path)
        
        if i > 0:
            merged_doc.add_page_break()
        
        for paragraph in doc.paragraphs:
            p = merged_doc.add_paragraph()
            for run in paragraph.runs:
                new_run = p.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
        
        for table in doc.tables:
            t = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    t.cell(i, j).text = cell.text
        
        os.remove(file_path)
    
    merged_filename = os.path.join(TEMP_FOLDER, f"{uuid.uuid4()}.docx")
    merged_doc.save(merged_filename)
    
    return merged_filename

def convert_to_markdown(docx_file):
    doc = Document(docx_file)
    
    markdown_content = ""
    
    h = html2text.HTML2Text()
    h.ignore_links = False
    h.ignore_images = False
    h.ignore_tables = False
    
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.replace('Heading ', ''))
            markdown_content += '#' * level + ' ' + paragraph.text + '\n\n'
        else:
            para_text = ""
            for run in paragraph.runs:
                text = run.text
                if run.bold and run.italic:
                    text = f"***{text}***"
                elif run.bold:
                    text = f"**{text}**"
                elif run.italic:
                    text = f"*{text}*"
                para_text += text
            
            markdown_content += para_text + "\n\n"
    
    for table in doc.tables:
        header_row = " | ".join(cell.text for cell in table.rows[0].cells)
        markdown_content += "| " + header_row + " |\n"
        markdown_content += "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |\n"
        
        for row in table.rows[1:]:
            row_text = " | ".join(cell.text for cell in row.cells)
            markdown_content += "| " + row_text + " |\n"
        
        markdown_content += "\n"
    
    os.remove(docx_file)
    
    return markdown_content.encode('utf-8'), 'merged.md', 'text/markdown'

def convert_to_text(docx_file):
    doc = Document(docx_file)
    
    text_content = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join([cell.text for cell in row.cells])
            text_content += "\n" + row_text
    
    os.remove(docx_file)
    
    return text_content.encode('utf-8'), 'merged.txt', 'text/plain'

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)